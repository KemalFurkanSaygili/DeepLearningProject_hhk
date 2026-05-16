"""
Generate the project report as a .docx file.
Run: python generate_report.py
Output: deep_learning_report.docx
"""

import os, glob
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ─────────────────────────────────────────────────────────────────────
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
found = glob.glob(os.path.join(
    os.path.expanduser('~'), 'OneDrive', '*', 'deep_learning_project', 'results'))
RESULTS_DIR = found[0] if found else os.path.join(SCRIPT_DIR, 'results')
OUT_PATH = os.path.join(SCRIPT_DIR, 'deep_learning_report.docx')

def img(name):
    return os.path.join(RESULTS_DIR, name)

# ── Document setup ────────────────────────────────────────────────────────────
doc = Document()

# Page margins (same as example: 2.5cm all sides)
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default paragraph font
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def set_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.runs[0]
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        run.font.size = Pt(14)
        run.bold = True
    else:
        run.font.size = Pt(12)
        run.bold = True
    return h

def add_para(doc, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold   = bold
    run.italic = italic
    return p

def add_figure(doc, img_path, caption, width=Inches(5.5)):
    if not os.path.exists(img_path):
        print(f'  WARNING: {img_path} not found, skipping.')
        return
    doc.add_picture(img_path, width=width)
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.name = 'Times New Roman'
    cap.runs[0].font.size = Pt(11)
    cap.runs[0].italic = True

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run('ESKISEHIR TECHNICAL UNIVERSITY\nFACULTY OF ENGINEERING\nDEPARTMENT OF COMPUTER ENGINEERING')
tr.font.name = 'Times New Roman'
tr.font.size = Pt(14)
tr.bold = True

doc.add_paragraph()

course_p = doc.add_paragraph()
course_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
cr = course_p.add_run('BIM 308 – DEEP LEARNING\nPROJECT REPORT')
cr.font.name = 'Times New Roman'
cr.font.size = Pt(16)
cr.bold = True

doc.add_paragraph()
doc.add_paragraph()

proj_p = doc.add_paragraph()
proj_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pr = proj_p.add_run('Steel Surface Defect Classification\nUsing Convolutional Neural Networks (CNN)\non NEU-DET Dataset')
pr.font.name = 'Times New Roman'
pr.font.size = Pt(14)
pr.bold = True

doc.add_paragraph()
doc.add_paragraph()

author_p = doc.add_paragraph()
author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
ar = author_p.add_run('Prepared by:\nKemal Furkan Saygılı')
ar.font.name = 'Times New Roman'
ar.font.size = Pt(12)

doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
dr = date_p.add_run('May 2025')
dr.font.name = 'Times New Roman'
dr.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. PROBLEM DEFINITION
# ═══════════════════════════════════════════════════════════════════════════════
set_heading(doc, '1. Problem Definition')

add_para(doc,
    'Steel surface defects are a critical quality control issue in the manufacturing industry. '
    'Identifying these defects manually is time-consuming, error-prone, and costly. '
    'Automated defect detection using deep learning offers a reliable and scalable alternative '
    'to traditional inspection methods.')

add_para(doc,
    'In this project, we address the problem of classifying six types of steel surface defects '
    'using Convolutional Neural Networks (CNN). The classification task involves assigning each '
    'grayscale steel surface image to exactly one of the following six defect categories:')

for cls in ['Crazing – a network of fine surface cracks',
            'Inclusion – embedded foreign particles within the steel',
            'Patches – irregular discolored areas on the surface',
            'Pitted Surface – small pits or holes on the surface',
            'Rolled-in Scale – scale pressed into the surface during rolling',
            'Scratches – linear grooves or scratches on the surface']:
    add_bullet(doc, cls)

add_para(doc,
    'The key challenges of this problem include the visual similarity between certain defect '
    'classes (e.g., scratches vs. pitted surface), relatively low image resolution (200×200 pixels), '
    'and the need for high accuracy to be applicable in industrial settings.')

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. NEURAL NETWORK MODEL EXPLANATION
# ═══════════════════════════════════════════════════════════════════════════════
set_heading(doc, '2. Neural Network Models Used')

add_para(doc,
    'Two CNN-based models were implemented and compared in this project: a Custom CNN built '
    'from scratch, and a pre-trained ResNet18 adapted via Transfer Learning.')

set_heading(doc, '2.1 Custom CNN Architecture', level=2)

add_para(doc,
    'The Custom CNN was designed specifically for this classification task. '
    'It consists of four convolutional blocks followed by global average pooling and '
    'a fully connected classification head:')

for item in [
    'Conv Block 1: 3 → 32 filters (Conv2D × 2, BatchNorm, ReLU, MaxPool 2×2) — output: 112×112',
    'Conv Block 2: 32 → 64 filters (Conv2D × 2, BatchNorm, ReLU, MaxPool 2×2) — output: 56×56',
    'Conv Block 3: 64 → 128 filters (Conv2D × 2, BatchNorm, ReLU, MaxPool 2×2) — output: 28×28',
    'Conv Block 4: 128 → 256 filters (Conv2D × 2, BatchNorm, ReLU, MaxPool 2×2) — output: 14×14',
    'Global Average Pooling → 256-dimensional feature vector',
    'Fully Connected: 256 → 512 → Dropout(0.5) → 6 (softmax output)',
]:
    add_bullet(doc, item)

add_para(doc,
    'Batch Normalization is applied after each convolutional layer to stabilize training '
    'and accelerate convergence. Dropout (p=0.5) is applied in the classifier head to '
    'reduce overfitting. The total number of trainable parameters is approximately 2.4 million.')

set_heading(doc, '2.2 ResNet18 with Transfer Learning', level=2)

add_para(doc,
    'ResNet18 is a 18-layer deep residual network pre-trained on the ImageNet dataset '
    '(1.2 million images, 1000 classes). Transfer learning allows us to leverage the '
    'powerful feature representations already learned from large-scale data.')

add_para(doc,
    'For this project, all layers of ResNet18 were frozen (weights not updated during training) '
    'and only the final classification head was replaced and trained:')

for item in [
    'Frozen backbone: ResNet18 (11,176,512 parameters — not updated)',
    'New classification head: FC(512→256) → ReLU → Dropout(0.4) → FC(256→6)',
    'Trainable parameters: 132,870 (~1.2% of total)',
]:
    add_bullet(doc, item)

add_para(doc,
    'This approach is computationally efficient and particularly effective when the target '
    'dataset is relatively small, as ImageNet features generalize well to texture-based '
    'classification tasks such as defect detection.')

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 3. DATASET
# ═══════════════════════════════════════════════════════════════════════════════
set_heading(doc, '3. Dataset Description')

add_para(doc,
    'The NEU-DET (NEU Surface Defect Database) dataset was used in this project. '
    'It is a widely used benchmark dataset for steel surface defect detection and '
    'classification, provided by Northeastern University, China.')

add_para(doc, 'Dataset statistics:')

table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
for i, txt in enumerate(['Split', 'Images per Class', 'Total Classes', 'Total Images']):
    hdr[i].text = txt
    hdr[i].paragraphs[0].runs[0].bold = True
    hdr[i].paragraphs[0].runs[0].font.name = 'Times New Roman'

rows_data = [
    ('Train',      '240', '6', '1,440'),
    ('Validation', '60',  '6', '360'),
    ('Total',      '300', '6', '1,800'),
]
for i, row_data in enumerate(rows_data):
    row = table.rows[i+1].cells
    for j, val in enumerate(row_data):
        row[j].text = val
        row[j].paragraphs[0].runs[0].font.name = 'Times New Roman'
        row[j].paragraphs[0].runs[0].font.size = Pt(11)

doc.add_paragraph()

add_para(doc,
    'Images are grayscale, 200×200 pixels, and organized into per-class subdirectories. '
    'The dataset is perfectly balanced with equal samples per class in both splits.')

add_para(doc,
    'Data augmentation was applied to the training set to improve generalization:')
for aug in ['Random horizontal and vertical flipping',
            'Random rotation (±15°)',
            'Color jitter (brightness, contrast, saturation)',
            'Resize to 224×224 pixels (required for ResNet18 input)']:
    add_bullet(doc, aug)

add_para(doc,
    'The dataset was obtained from the official NEU-DET distribution. '
    'It is a standardized benchmark frequently cited in industrial defect detection research.')

add_figure(doc, img('dataset_distribution.png'),
           'Figure 1. Class distribution of the NEU-DET dataset across train and validation splits.')

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. IMPLEMENTATION
# ═══════════════════════════════════════════════════════════════════════════════
set_heading(doc, '4. Implementation Details')

add_para(doc, 'The project was implemented in Python using the PyTorch deep learning framework. '
    'All experiments were run on CPU (Intel). The following hyperparameters were used for both models:')

for item in [
    'Optimizer: Adam (lr=1e-3, weight_decay=1e-4)',
    'Loss function: Cross-Entropy Loss',
    'Scheduler: ReduceLROnPlateau (factor=0.5, patience=3)',
    'Batch size: 32',
    'Number of epochs: 20',
    'Image size: 224×224 (RGB)',
]:
    add_bullet(doc, item)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. SIMULATION RESULTS
# ═══════════════════════════════════════════════════════════════════════════════
set_heading(doc, '5. Simulation Results')

add_para(doc,
    'Both models were trained for 20 epochs and evaluated on the validation set. '
    'The best model state (highest validation accuracy) was saved and used for final evaluation.')

set_heading(doc, '5.1 Training History', level=2)
add_figure(doc, img('training_history.png'),
           'Figure 2. Training and validation loss and accuracy curves for both models over 20 epochs.')

set_heading(doc, '5.2 Custom CNN Results', level=2)

add_para(doc, 'The Custom CNN achieved the following results on the validation set:')
for item in ['Validation Accuracy: 97.78%',
             'Macro F1-Score: 0.9777',
             'Best performance on: crazing, patches, rolled-in_scale (F1 = 1.00)']:
    add_bullet(doc, item)

add_figure(doc, img('cm_custom_cnn.png'),
           'Figure 3. Confusion matrix of the Custom CNN on the validation set.')

add_figure(doc, img('samples_custom_cnn.png'),
           'Figure 4. Sample predictions of the Custom CNN (green = correct, red = incorrect).')

set_heading(doc, '5.3 ResNet18 Transfer Learning Results', level=2)

add_para(doc, 'The ResNet18 model with Transfer Learning achieved the following results:')
for item in ['Validation Accuracy: 98.61%',
             'Macro F1-Score: 0.9861',
             'Correctly classified: 355 out of 360 images']:
    add_bullet(doc, item)

add_figure(doc, img('cm_resnet18.png'),
           'Figure 5. Confusion matrix of ResNet18 (Transfer Learning) on the validation set.')

add_figure(doc, img('samples_resnet18.png'),
           'Figure 6. Sample predictions of ResNet18 (green = correct, red = incorrect).')

set_heading(doc, '5.4 Model Comparison', level=2)

add_figure(doc, img('model_comparison.png'),
           'Figure 7. Comparison of validation accuracy and macro F1-score for both models.')

# Summary table
add_para(doc, 'Summary of results:')
table2 = doc.add_table(rows=3, cols=4)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
for i, txt in enumerate(['Model', 'Val Accuracy', 'Macro F1', 'Trainable Params']):
    hdr2[i].text = txt
    hdr2[i].paragraphs[0].runs[0].bold = True
    hdr2[i].paragraphs[0].runs[0].font.name = 'Times New Roman'

rows_data2 = [
    ('Custom CNN',         '97.78%', '0.9777', '~2.4M'),
    ('ResNet18 (TL)',      '98.61%', '0.9861', '132,870'),
]
for i, row_data in enumerate(rows_data2):
    row = table2.rows[i+1].cells
    for j, val in enumerate(row_data):
        row[j].text = val
        row[j].paragraphs[0].runs[0].font.name = 'Times New Roman'
        row[j].paragraphs[0].runs[0].font.size = Pt(11)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 6. DISCUSSION & ANALYSIS
# ═══════════════════════════════════════════════════════════════════════════════
set_heading(doc, '6. Discussion and Analysis')

add_para(doc,
    'Both models achieved excellent performance on the NEU-DET validation set, '
    'with accuracies above 97%. This demonstrates that CNN-based architectures are '
    'well-suited for steel surface defect classification.')

add_para(doc,
    'Transfer Learning (ResNet18) outperformed the Custom CNN by approximately 0.83 percentage '
    'points in accuracy. Despite using only 132,870 trainable parameters (compared to ~2.4M for '
    'Custom CNN), ResNet18 benefited from rich feature representations pre-trained on ImageNet, '
    'enabling faster convergence and better generalization.')

add_para(doc,
    'The confusion matrices reveal that the most challenging class pair is '
    'pitted_surface vs. scratches. These two defects share similar low-level '
    'texture features (sharp discontinuities on the surface), making them harder '
    'to distinguish even for deep networks.')

add_para(doc,
    'The training curves show that both models converged smoothly without significant '
    'overfitting. The ReduceLROnPlateau scheduler effectively decreased the learning '
    'rate when validation loss plateaued, helping models find better optima.')

add_para(doc,
    'Data augmentation (flipping, rotation, color jitter) played a key role in '
    'preventing overfitting on the relatively small training set of 1,440 images.')

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 7. CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════════
set_heading(doc, '7. Conclusion')

add_para(doc,
    'In this project, we successfully applied Convolutional Neural Networks to the task '
    'of steel surface defect classification using the NEU-DET dataset. Two approaches '
    'were implemented and compared:')

for item in [
    'Custom CNN trained from scratch: 97.78% validation accuracy',
    'ResNet18 with Transfer Learning: 98.61% validation accuracy',
]:
    add_bullet(doc, item)

add_para(doc,
    'Both models achieved strong results, with the Transfer Learning approach proving '
    'slightly superior. The results confirm that deep learning is an effective and '
    'practical solution for automated quality control in steel manufacturing.')

add_para(doc,
    'Future work could explore: (1) fine-tuning all ResNet18 layers with a lower learning '
    'rate, (2) using more advanced architectures such as EfficientNet or Vision Transformers, '
    '(3) multi-label classification for images containing multiple defect types, and '
    '(4) deploying the model as a real-time inspection system.')

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 8. HOW TO RUN
# ═══════════════════════════════════════════════════════════════════════════════
set_heading(doc, '8. How to Run the Code')

set_heading(doc, 'Prerequisites', level=2)
add_para(doc, 'Ensure you have Python 3.8+ installed.')
add_para(doc, 'Install the required libraries:')

code_p = doc.add_paragraph()
run = code_p.add_run(
    'python -m pip install torch torchvision matplotlib scikit-learn seaborn pdfplumber')
run.font.name = 'Courier New'
run.font.size = Pt(10)

set_heading(doc, 'Directory Setup', level=2)
add_para(doc, 'Place the following files in the same directory:')
for item in ['deeplearning.py  (main training script)',
             'NEU-DET/  (dataset folder with train/ and validation/ subdirectories)',
             'test.py  (optional: for standalone evaluation)']:
    add_bullet(doc, item)

set_heading(doc, 'Training', level=2)
code2 = doc.add_paragraph()
run2 = code2.add_run('python deeplearning.py')
run2.font.name = 'Courier New'
run2.font.size = Pt(10)

add_para(doc,
    'Training will run for 20 epochs. Results (plots and model weights) will be '
    'saved automatically to the results/ folder.')

set_heading(doc, 'Submission Mode (100 samples)', level=2)
add_para(doc, 'To run with only 100 dataset samples, open deeplearning.py and set:')
code3 = doc.add_paragraph()
run3 = code3.add_run('SUBMISSION_MODE = True')
run3.font.name = 'Courier New'
run3.font.size = Pt(10)
add_para(doc, 'Then run: python deeplearning.py')

set_heading(doc, 'Evaluation', level=2)
code4 = doc.add_paragraph()
run4 = code4.add_run(
    'python test.py --model resnet18\n'
    'python test.py --image "NEU-DET/validation/images/crazing/crazing_241.jpg"')
run4.font.name = 'Courier New'
run4.font.size = Pt(10)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════════════════════
set_heading(doc, 'References')

refs = [
    'K. Song and Y. Yan, "A noise robust method based on completed local binary patterns for hot-rolled steel strip surface defects," Applied Surface Science, vol. 285, pp. 858–864, 2013.',
    'He, K., Zhang, X., Ren, S., & Sun, J. (2016). Deep residual learning for image recognition. In Proceedings of the IEEE conference on computer vision and pattern recognition (pp. 770–778).',
    'NEU Surface Defect Database. http://faculty.neu.edu.cn/yunhyan/NEU_surface_defect_database.html',
    'Paszke, A., et al. (2019). PyTorch: An imperative style, high-performance deep learning library. Advances in Neural Information Processing Systems, 32.',
    'Pan, S. J., & Yang, Q. (2009). A survey on transfer learning. IEEE Transactions on knowledge and data engineering, 22(10), 1345–1359.',
    'Litjens, G., et al. (2017). A survey on deep learning in medical image analysis. Medical image analysis, 42, 60–88.',
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT_PATH)
print(f'\nReport saved: {OUT_PATH}')
